from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
OUT_DOWNLOADS = Path.home() / "Downloads" / "Population-Insight_8项功能实现与运行验证报告.docx"
OUT_PROJECT = BASE / "output" / "Population-Insight_8项功能实现与运行验证报告.docx"

IMAGES = {
    "dashboard": BASE / "output/playwright/desktop-dashboard.png",
    "records": BASE / "output/playwright/desktop-records.png",
    "add_form": BASE / "output/mysql_evidence/02_platform_add_form_filled.png",
    "query_result": BASE / "output/mysql_evidence/04_platform_query_result.png",
    "collection": BASE / "output/doc_original_images/original_07.png",
    "statistics": BASE / "output/doc_original_images/original_06.png",
    "charts": BASE / "output/playwright/desktop-charts.png",
    "db_population": BASE / "output/table_evidence/db_population_data.png",
}

BLACK = RGBColor(0, 0, 0)


def ensure_required_images():
    missing = [str(path) for path in IMAGES.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("\n".join(missing))


def set_run_font(run, size=None, bold=None):
    run.font.name = "Times New Roman"
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:cs"), "Times New Roman")


def set_style_font(style, size=None):
    style.font.name = "Times New Roman"
    style.font.color.rgb = BLACK
    if size is not None:
        style.font.size = Pt(size)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:cs"), "Times New Roman")


def normalize_fonts(doc):
    for style in doc.styles:
        if style.type in (1, 2):
            try:
                set_style_font(style)
            except Exception:
                pass
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            size = run.font.size.pt if run.font.size is not None else None
            set_run_font(run, size=size, bold=run.bold)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        size = run.font.size.pt if run.font.size is not None else None
                        set_run_font(run, size=size, bold=run.bold)


def add_paragraph(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.runs[0]
    set_run_font(r, size=16 if level == 1 else 13, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=10)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    return p


def add_image(doc, image_key, caption, width=5.8):
    path = IMAGES[image_key]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_summary_table(doc):
    rows = [
        ("新增", "已实现", "record_create_view() / add_population_record()", "人口数据管理页面"),
        ("删除", "已实现", "record_delete_view() / delete_population_record()", "表格操作列删除按钮"),
        ("修改", "已实现", "record_edit_view() / update_population_record()", "表格操作列编辑按钮"),
        ("查询", "已实现", "records_view() / query_population_records()", "多条件筛选区域"),
        ("批量导入", "已实现", "collection_view() / import_collected_records()", "人口数据采集页面"),
        ("CSV 导出", "已实现", "export_records() / export_to_csv()", "导出当前结果按钮"),
        ("数据统计", "已实现", "statistics_view() / calculate_statistics()", "统计分析页面"),
        ("可视化绘图", "已实现", "chart_trend_api() / chart_bar_api()", "图表中心页面"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["检查功能", "实现情况", "核心函数", "对应页面/操作"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_feature_section(doc, idx, title, desc, steps, funcs, image_key, caption):
    add_heading(doc, f"{idx}. {title}", level=2)
    add_paragraph(doc, desc)
    add_paragraph(doc, "实现流程：", bold_prefix="实现流程：")
    for step in steps:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        r = p.add_run(step)
        set_run_font(r)
    add_paragraph(doc, f"核心函数：{funcs}", bold_prefix="核心函数：")
    add_image(doc, image_key, caption)


def add_code_block(doc, code):
    p = doc.add_paragraph()
    for line in code.strip().splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = BLACK
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Consolas")
        rfonts.set(qn("w:hAnsi"), "Consolas")
        rfonts.set(qn("w:eastAsia"), "宋体")


def build_document():
    ensure_required_images()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Population-Insight 8项功能实现与运行验证报告")
    set_run_font(run, size=18, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("对应检查标准：8 项功能实现、Web 页面可交互、数据库读取到表格展示与图表绘制")
    set_run_font(run, size=11)

    add_heading(doc, "一、系统功能概述")
    add_paragraph(
        doc,
        "本系统基于 Web 页面实现人口数据管理与分析功能，后端使用 Flask 提供页面路由和接口服务，"
        "数据存储在数据库中。系统支持人口数据的增、删、改、查、批量导入、CSV 导出、数据统计和可视化绘图。"
        "用户可通过网页按钮、表单和筛选条件完成交互操作，系统能够从数据库读取数据后展示到页面表格，"
        "并进一步绘制折线图或柱状图。",
    )

    add_heading(doc, "二、检查标准对应情况")
    add_summary_table(doc)

    add_heading(doc, "三、Web 页面交互验证")
    add_paragraph(
        doc,
        "本项目采用 Web 页面作为交互界面。用户登录后可通过左侧导航进入数据管理、数据采集、统计分析、"
        "图表中心等页面，页面中的按钮、输入框、下拉框和表格均可触发对应后端功能。",
    )
    add_image(doc, "dashboard", "图1-1 系统 Web 概览页面")
    add_image(doc, "records", "图1-2 人口数据管理页面及交互按钮")

    add_heading(doc, "四、八项功能实现说明")
    features = [
        (
            "新增功能",
            "管理员在“人口数据管理”页面点击“新增记录”，填写地区、年份、总人口、出生率等字段，提交后调用 add_population_record() 写入数据库。",
            ["点击新增记录按钮", "填写表单字段", "后端校验并写入 population_data 表", "返回列表页展示新增数据"],
            "record_create_view() / add_population_record()",
            "add_form",
            "图1-3 新增人口数据表单",
        ),
        (
            "删除功能",
            "用户在数据列表中选择目标记录，点击删除按钮，系统根据 record_id 调用 delete_population_record() 删除数据库中的对应记录。",
            ["用户点击表格中的删除按钮", "后端接收 record_id", "执行 DELETE SQL", "页面刷新后不再显示该记录"],
            "record_delete_view() / delete_population_record()",
            "records",
            "图1-4 数据列表中的删除操作入口",
        ),
        (
            "修改功能",
            "系统支持对已有记录进行编辑。点击编辑后页面回填原数据，用户修改后提交，后端调用 update_population_record() 更新数据库。",
            ["点击编辑按钮", "根据 ID 读取原记录", "修改表单并提交", "执行 UPDATE SQL 更新数据"],
            "record_edit_view() / update_population_record()",
            "records",
            "图1-5 数据列表中的编辑操作入口",
        ),
        (
            "查询功能",
            "数据管理页面支持按地区、年份、起止年份、人口规模、出生率等多条件筛选数据。后端将请求参数解析为 filters，再调用 query_population_records() 查询数据库。",
            ["输入筛选条件", "解析 request.args", "组装 WHERE 条件", "查询结果在表格展示"],
            "records_view() / query_population_records()",
            "query_result",
            "图1-6 多条件查询后的数据结果",
        ),
        (
            "批量导入功能",
            "数据采集页面支持粘贴表格文本或上传 CSV 文件。系统先解析数据并给出预览，用户确认后批量导入数据库。",
            ["上传 CSV 或粘贴表格文本", "调用 parse_population_records()", "预览解析结果", "确认后调用 import_collected_records() 批量入库"],
            "collection_view() / import_collected_records()",
            "collection",
            "图1-7 人口数据采集与批量导入页面",
        ),
        (
            "CSV 导出功能",
            "用户可将当前查询或筛选结果导出为 CSV 文件。后端复用查询条件，调用 export_to_csv() 生成文件，再通过浏览器下载。",
            ["用户点击“导出当前结果”", "后端获取当前 filters", "调用 export_to_csv()", "返回 CSV 文件下载响应"],
            "export_records() / export_to_csv()",
            "records",
            "图1-8 CSV 导出按钮与数据管理页面",
        ),
        (
            "数据统计功能",
            "统计分析页面按地区、年份范围和指标计算统计结果，包括记录数、最大值、最小值、平均值和地区排名等。",
            ["选择统计条件", "调用 calculate_statistics()", "从数据库读取数据", "在页面输出统计卡片和排名表"],
            "statistics_view() / calculate_statistics()",
            "statistics",
            "图1-9 数据统计分析页面",
        ),
        (
            "可视化绘图功能",
            "图表中心页面从数据库读取人口年度数据，将数据组装为 JSON，前端使用 ECharts 绘制趋势折线图或柱状图。",
            ["选择地区和指标", "前端请求 chart API", "后端查询数据库并返回 JSON", "前端绘制折线图"],
            "chart_trend_api() / chart_bar_api()",
            "charts",
            "图1-10 图表中心中的人口趋势折线图",
        ),
    ]
    for idx, feature in enumerate(features, start=1):
        add_feature_section(doc, idx, *feature)

    add_heading(doc, "五、核心代码片段")
    add_code_block(
        doc,
        '''
@app.route("/records/new", methods=["GET", "POST"])
def record_create_view():
    if request.method == "POST":
        record_id = add_population_record(request.form.to_dict(), g.user["username"])
        return redirect(url_for("records_view"))
    return render_template("record_form.html")

@app.route("/records/<int:record_id>/delete", methods=["POST"])
def record_delete_view(record_id: int):
    delete_population_record(record_id, g.user["username"])
    return redirect(url_for("records_view"))

@app.route("/export")
def export_records():
    filters, _ = _parse_record_filters(request.args)
    records = query_population_records(filters)
    csv_path = export_to_csv(records)
    return send_file(csv_path, as_attachment=True)

@app.route("/api/charts/trend")
def chart_trend_api():
    data = get_chart_trend_data_with_mode(region, metric, mode)
    return jsonify({"success": True, "data": data})
''',
    )

    add_heading(doc, "六、数据库读取到表格展示")
    add_paragraph(
        doc,
        "系统中的人口数据统一存储在 population_data 表中。页面展示数据时，后端先调用 query_population_records() "
        "读取数据库记录，再将 records 传递给 records.html 模板，最终在网页表格中展示。",
    )
    for item in [
        "population_data 数据表",
        "query_population_records() 执行查询",
        "records_view() 接收 records",
        "records.html 渲染表格",
    ]:
        p = doc.add_paragraph(style=doc.styles["List Bullet"])
        set_run_font(p.add_run(item))
    add_image(doc, "db_population", "图1-11 数据库 population_data 表中的人口数据")
    add_image(doc, "query_result", "图1-12 数据库记录在 Web 表格中的展示")

    add_heading(doc, "七、数据库读取到图表绘制")
    add_paragraph(
        doc,
        "图表绘制时，前端页面向后端 API 发送地区、年份和指标参数。后端从数据库读取对应人口数据，"
        "整理为 labels、values 等 JSON 结构返回给前端，前端使用 ECharts 绘制折线图或柱状图。",
    )
    for item in ["数据库记录 → 后端查询函数", "后端 API → JSON 图表数据", "前端 ECharts → 折线图/柱状图"]:
        p = doc.add_paragraph(style=doc.styles["List Bullet"])
        set_run_font(p.add_run(item))
    add_image(doc, "charts", "图1-13 由数据库数据生成的人口趋势折线图")

    add_heading(doc, "八、总结")
    add_paragraph(
        doc,
        "综上，本系统已完成新增、删除、修改、查询、批量导入、CSV 导出、数据统计和可视化绘图 8 项核心功能。"
        "系统采用 Web 页面实现交互，用户能够通过按钮、表单和筛选条件触发功能。人口数据从数据库读取后能够在表格中展示，"
        "并能进一步生成统计结果和可视化图表，满足检查标准中关于功能完整性、界面可用性和图表展示的要求。",
    )

    normalize_fonts(doc)
    OUT_PROJECT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PROJECT)
    OUT_DOWNLOADS.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUT_PROJECT, OUT_DOWNLOADS)
    print(f"downloads={OUT_DOWNLOADS}")
    print(f"project={OUT_PROJECT}")
    print(f"images={len(doc.inline_shapes)}")


if __name__ == "__main__":
    build_document()
