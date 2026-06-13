from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Population-Insight整体开发思路文档.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C4D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), BORDER)


def set_table_geometry(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_borders(table)


def set_east_asian_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc: Document, text: str = "", style: str = "Normal", bold_prefix: str | None = None):
    para = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = para.add_run(bold_prefix)
        set_east_asian_font(run, "Microsoft YaHei")
        run.bold = True
        run.font.size = Pt(11)
        para.add_run(text[len(bold_prefix):])
    else:
        para.add_run(text)
    for run in para.runs:
        set_east_asian_font(run, "Microsoft YaHei")
        if run.font.size is None:
            run.font.size = Pt(11)
    return para


def add_heading(doc: Document, text: str, level: int = 1):
    para = doc.add_paragraph(style=f"Heading {level}")
    run = para.add_run(text)
    set_east_asian_font(run, "Microsoft YaHei")
    run.bold = True
    return para


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        set_east_asian_font(run, "Microsoft YaHei")
        run.font.size = Pt(11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(item)
        set_east_asian_font(run, "Microsoft YaHei")
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_GRAY)
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_east_asian_font(run, "Microsoft YaHei")
        run.bold = True
        run.font.size = Pt(10.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_east_asian_font(run, "Microsoft YaHei")
            run.font.size = Pt(9.5)
    set_table_geometry(table, widths_cm)
    doc.add_paragraph()
    return table


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color, before, after in [
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Population-Insight 人口趋势与分析管理系统"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_east_asian_font(run, "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)
    footer = section.footer.paragraphs[0]
    footer.text = "整体开发思路文档"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_east_asian_font(run, "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("项目整体开发思路文档")
    set_east_asian_font(run, "Microsoft YaHei")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Population-Insight 人口趋势与分析管理系统")
    set_east_asian_font(run, "Microsoft YaHei")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(71, 85, 105)

    add_heading(doc, "一．项目概述", 1)
    add_paragraph(
        doc,
        "本项目基于 Python + Flask + SQLite + ECharts + Matplotlib 实现人口趋势预测、人口数据采集管理、统计排名、风险预警、分析报告和 CSV 导出一体化系统。系统围绕年度人口核心指标建立统一数据底座，通过 Web 工作台和控制台入口实现数据治理、可视化分析、机器学习预测和成果输出闭环。"
    )
    add_paragraph(
        doc,
        "项目当前预置数据库覆盖 31 个地区、2015-2024 年共 310 条人口年度记录，并包含用户、地区档案、数据来源、扩展指标、年度扩展指标、分析报告和操作日志等支撑数据。整体架构分为配置与数据层、数据库访问层、业务服务层、Web 应用层、前端可视化层、系统输出与文档层六大模块。"
    )

    add_heading(doc, "二．主要功能", 1)
    add_numbered(
        doc,
        [
            "机器学习人口趋势预测：基于历史人口数据进行线性回归、复合增长率拟合、近期趋势拟合和融合预测。",
            "大模型结果解释分析：支持 OpenAI 兼容接口输出趋势判断、风险原因和政策建议；未配置接口时自动使用本地规则解释。",
            "智能风险评估：依据老龄化率、出生率、自然增长率和预测期人口变化计算风险得分与风险等级。",
            "数据可视化分析：提供趋势图、柱状图、性别结构饼图、预测结果图和仪表盘指标展示。",
            "统计分析与地区排名：统计平均人口、平均出生率、平均老龄化率、最大最小人口地区和地区增长排名。",
            "人口数据采集：支持公开 URL、HTML 表格、CSV、制表符文本和粘贴表格文本解析导入。",
            "人口数据管理：支持人口年度数据新增、修改、删除、查询、分页和操作日志记录。",
            "分析报告生成：支持将筛选条件、统计摘要、预测解释和人工摘要保存为分析报告。",
            "多条件查询筛选：支持地区、年份、年份区间、总人口区间、出生率区间、老龄化率区间组合查询。",
            "数据排序与对比：支持总人口、年份、出生率、老龄化率等字段排序，并提供多地区指标横向对比。",
            "地区档案管理：维护地区名称、地区类型、行政区划代码、上级地区和备注。",
            "数据来源管理：维护来源名称、发布机构、来源链接、发布日期、可信等级和备注。",
            "扩展指标管理：维护指标编码、指标名称、单位、说明，并支持年度扩展指标值录入。",
            "CSV 数据导出：按当前筛选结果导出 UTF-8 BOM CSV 文件，便于离线分析和答辩展示。",
        ],
    )

    add_heading(doc, "三．分层架构设计", 1)
    add_heading(doc, "（1）配置与基础数据层：config.py + data/population_insight.db", 2)
    add_paragraph(
        doc,
        "config.py 统一维护数据库路径、输出目录、默认用户、指标标签、排序白名单、分页大小和初始化样例数据；data/population_insight.db 作为 SQLite 数据库保存全部业务数据。系统首次启动时自动创建 data、output/charts、output/exports 目录，完成建表、默认用户写入、人口样例数据导入和扩展数据初始化。"
    )

    add_heading(doc, "（2）数据库访问层：db/connection.py + db/initializer.py", 2)
    add_paragraph(
        doc,
        "connection.py 封装 get_connection、fetch_one、fetch_all、execute_write 等通用数据库操作，统一 SQLite row_factory、提交和连接管理；initializer.py 负责创建 users、population_data、operation_logs、regions、data_sources、population_indicators、annual_indicator_values、analysis_reports 八张核心表，并通过唯一约束、CHECK 约束和外键约束保证基础数据质量。"
    )

    add_heading(doc, "（3）业务服务层：population_insight/services", 2)
    add_paragraph(doc, "模块化拆分业务逻辑，解耦 Web 路由和控制台入口，app.py 与 main.py 共用同一套服务函数，避免重复实现。")
    add_bullets(
        doc,
        [
            "population_service：人口数据新增、修改、删除、查询、排序、地区年份枚举和日志联动。",
            "collection_service：公开 URL 拉取、HTML 表格解析、CSV/文本解析、字段别名识别、单位换算和采集导入。",
            "statistics_service：条件统计、平均指标、最大最小人口记录、地区增长率汇总和地区排名。",
            "comparison_service：多地区多年份指标对比、折线图数据、对比表格、变化幅度和最新排名。",
            "dashboard_service：仪表盘汇总、数据健康度、最新年份热点、老龄化热点和图表接口数据。",
            "prediction_service：线性回归、复合增长率、近期趋势拟合、融合预测、结构变化识别、风险评估和解释分析。",
            "visualization_service：Matplotlib 趋势图、柱状图、性别结构饼图后端图片生成。",
            "extension_service：地区档案、数据来源、扩展指标、年度指标值、分析报告和人口预警。",
            "export_service：按查询结果生成 CSV 导出文件。",
            "auth_service / log_service：登录校验、角色识别、操作日志写入和日志查询。",
        ],
    )
    add_paragraph(doc, "设计优势：服务层函数同时服务 Web 页面、API 接口和命令行菜单，数据校验、日志记录、排序规则、预测逻辑只维护一份，便于测试、扩展和答辩讲解。")

    add_heading(doc, "（4）Web 应用层：app.py + web/auth.py", 2)
    add_bullets(
        doc,
        [
            "后端页面路由：登录、仪表盘、数据管理、数据采集、统计分析、地区对比、趋势预测、图表中心、预警分析、分析报告、地区档案、数据来源、扩展指标和操作日志。",
            "接口路由：图表数据、预测结果、地区对比等页面异步请求接口，前端根据 JSON 数据渲染 ECharts。",
            "权限控制：login_required 保证登录访问，admin_required 限制管理员维护、删除、采集和扩展管理操作。",
            "表单处理：统一解析筛选条件、排序字段、分页参数、年份和非负数校验，错误信息通过 flash 返回页面。",
        ],
    )

    add_heading(doc, "（5）前端可视化层：templates + static", 2)
    add_paragraph(
        doc,
        "templates 目录使用 Jinja2 组织页面结构，base.html 提供侧边栏导航和统一布局；records、collection、statistics、comparison、prediction、charts 等模板承载核心业务页面。static/css/app.css 统一界面样式，static/js 中的 charts.js、comparison.js、prediction.js、app.js 负责图表渲染、页面交互和动态数据展示。"
    )

    add_heading(doc, "（6）控制台与成果输出层：main.py + output + docs", 2)
    add_paragraph(
        doc,
        "main.py 提供控制台版本菜单，便于不启动 Web 时进行本地数据维护、查询、统计、导出和扩展管理；output/charts 存放后端生成图表，output/exports 存放 CSV 导出结果；docs 目录保存 ER/DFD、答辩材料和答辩 PPT，支撑课程设计提交与演示。"
    )

    add_heading(doc, "四．核心业务流程", 1)
    add_heading(doc, "（1）初始化部署流程", 2)
    add_numbered(
        doc,
        [
            "安装依赖：执行 pip install -r requirements.txt。",
            "启动 Web 系统：执行 python app.py，默认访问 http://127.0.0.1:5000。",
            "系统调用 init_database 自动创建目录、数据表和默认账号。",
            "写入 admin/admin123 管理员和 viewer/viewer123 普通用户。",
            "导入人口样例数据、地区档案、数据来源、扩展指标和年度扩展指标值。",
            "管理员登录后可继续补充人口数据、地区档案、数据来源和扩展指标。",
        ],
    )

    add_heading(doc, "（2）日常数据治理流程", 2)
    add_numbered(
        doc,
        [
            "管理员进入数据采集页面，输入公开统计数据 URL 或粘贴表格文本。",
            "collection_service 识别 HTML 表格、CSV 或分隔文本，匹配地区、年份、总人口、性别人口、出生率、死亡率、老龄化率、城镇化率等字段。",
            "系统完成单位换算、年份识别、非负数校验和采集预览。",
            "确认导入后写入 population_data，并记录 ADD_RECORD 或采集导入相关操作日志。",
            "在数据管理页面进行多条件筛选、排序、分页查看、编辑、删除和 CSV 导出。",
            "在地区档案和数据来源页面维护基础元数据，形成可追溯的数据来源说明。",
        ],
    )

    add_heading(doc, "（3）统计分析与可视化流程", 2)
    add_numbered(
        doc,
        [
            "用户选择地区、年份区间或指标条件，statistics_service 查询 population_data。",
            "系统计算记录数、平均总人口、平均出生率、平均老龄化率、最大最小人口地区和地区增长率。",
            "地区排名模块按总人口、出生率、老龄化率、城镇化率等指标生成 TOP 排名。",
            "图表中心调用 dashboard_service 或 visualization_service 生成趋势图、柱状图、性别结构饼图。",
            "地区对比页面调用 comparison_service 输出多地区年度序列、对比表、变化幅度和最新排名。",
        ],
    )

    add_heading(doc, "（4）预测、解释与风险评估流程", 2)
    add_numbered(
        doc,
        [
            "用户在趋势预测页选择地区、预测指标和预测年数，系统至少读取 3 年历史数据。",
            "prediction_service 分别计算线性回归、复合增长率和近期增长率，并按 0.5、0.25、0.25 权重融合预测结果。",
            "结构识别模块比较基期与最新年份的性别比例、老龄化率、出生率和城镇化率变化。",
            "风险评估模块根据老龄化率高位、出生率偏低、自然增长率为负和预测人口变化计算 0-100 分风险得分。",
            "解释模块优先调用 OpenAI 兼容大模型接口，未配置 API 时使用本地规则生成管理者可读的趋势说明和建议。",
            "预测图表和解释结果可保存为 analysis_reports，形成可回看、可展示的分析成果。",
        ],
    )

    add_heading(doc, "（5）系统支撑与成果输出流程", 2)
    add_bullets(
        doc,
        [
            "用户登录后系统通过 session 保存当前用户，页面全局注入当前用户和数据健康度。",
            "管理员维护地区、来源、扩展指标、年度指标值时均写入 operation_logs。",
            "预警页面集中展示老龄化率 >= 20%、自然增长率 < 0、出生率 < 7% 的风险记录。",
            "数据管理页可按当前筛选条件导出 CSV，导出文件保存到 output/exports。",
            "README、ER_DFD、DEFENSE_MATERIALS 和答辩 PPT 构成项目交付说明材料。",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "五．任务分工", 1)
    add_table(
        doc,
        ["成员", "负责模块", "主要工作内容"],
        [
            ["王景冉", "机器学习预测与智能决策分析模块", "负责基于历史人口数据的人口趋势预测、时间序列建模、线性回归预测、增长率拟合分析、人口结构变化识别、老龄化风险智能评估、预测结果大模型解释分析等功能。"],
            ["孙屿丰", "数据可视化与统计分析模块", "负责多条件筛选、数据排序、统计分析、地区排名、趋势图、柱状图、性别结构饼图、预测结果图表展示等功能。"],
            ["张金涛", "数据采集与基础数据管理模块", "负责人口数据爬取、公开统计数据采集、数据来源整理、人口数据新增、修改、删除、查询、地区档案管理等功能。"],
            ["密政奇", "系统支撑与成果输出模块", "负责用户登录与权限验证、扩展指标管理、年度扩展指标录入、CSV 数据导出、操作日志、分析报告、人口预警、README 文档、ER 图、DFD 图和答辩材料整理。"],
        ],
        [2.2, 4.0, 10.3],
    )

    doc.add_page_break()
    add_heading(doc, "六．数据关系表", 1)
    add_paragraph(
        doc,
        "项目数据库由 initializer.py 创建 8 张核心表。当前实际数据库包含 users 2 条、population_data 310 条、operation_logs 21 条、regions 31 条、data_sources 31 条、population_indicators 3 条、annual_indicator_values 6 条、analysis_reports 1 条记录。"
    )
    add_table(
        doc,
        ["表名", "核心字段", "作用"],
        [
            ["users", "id、username、password_hash、role、created_at", "存储用户账号、密码哈希和角色，支持管理员与普通用户权限区分。"],
            ["population_data", "id、region、year、total_population、male_population、female_population、birth_rate、death_rate、natural_growth_rate、aging_rate、urbanization_rate、data_source_name、source_url、data_quality", "人口年度核心数据表，承担查询、统计、对比、可视化、预测和导出的主要数据来源；region + year 唯一。"],
            ["operation_logs", "id、username、action、target_id、details、action_time", "记录新增、修改、删除、采集、扩展维护、报告生成等操作，实现审计追踪。"],
            ["regions", "id、name、region_type、admin_code、parent_region、remarks", "维护地区档案，与人口年度数据和扩展指标值通过地区名称形成逻辑对应。"],
            ["data_sources", "id、name、publisher、source_url、published_date、reliability_level、remarks", "维护公开统计数据来源，支撑人口数据的数据来源说明和可信度管理。"],
            ["population_indicators", "id、code、name、unit、description", "维护可扩展人口指标定义，code 唯一。"],
            ["annual_indicator_values", "id、region、year、indicator_code、value、remarks", "记录某地区、某年份、某扩展指标的具体值；indicator_code 外键关联 population_indicators.code，region + year + indicator_code 唯一。"],
            ["analysis_reports", "id、title、username、filter_summary、report_summary、created_at", "保存统计分析、预测解释和人工摘要，形成可回看的分析报告。"],
        ],
        [3.0, 6.0, 7.5],
    )

    add_heading(doc, "（1）表间关系说明", 2)
    add_bullets(
        doc,
        [
            "users.username 与 operation_logs.username、analysis_reports.username 形成逻辑关联，用于追踪用户行为和报告归属。",
            "regions.name 与 population_data.region、annual_indicator_values.region 形成逻辑关联，用于地区档案和年度数据对应。",
            "population_indicators.code 与 annual_indicator_values.indicator_code 是外键关系，用于扩展指标定义和值的对应。",
            "population_data 是统计分析、地区对比、图表展示、趋势预测、风险预警和 CSV 导出的核心事实表。",
            "data_sources 用于记录公开统计数据来源，与 population_data 中的 data_source_name、source_url 保持业务来源关联。",
            "analysis_reports 保存由 population_data 过滤统计和 prediction_service 预测解释形成的分析成果。",
        ],
    )

    add_heading(doc, "（2）核心约束设计", 2)
    add_bullets(
        doc,
        [
            "users.username、regions.name、data_sources.name、population_indicators.code 均设置唯一约束，避免重复基础资料。",
            "population_data 设置 UNIQUE(region, year)，保证同一地区同一年份只有一条人口年度数据。",
            "annual_indicator_values 设置 UNIQUE(region, year, indicator_code)，保证扩展指标值不重复录入。",
            "population_data 对总人口、男女人口、出生率、死亡率、老龄化率和城镇化率设置非负 CHECK 约束。",
            "users.role 设置 CHECK(role IN ('admin', 'viewer'))，保障角色值合法。",
        ],
    )

    add_heading(doc, "七．技术选型说明", 1)
    add_bullets(
        doc,
        [
            "后端语言：Python 3，便于数据处理、预测建模和课程项目快速实现。",
            "Web 框架：Flask，路由清晰、依赖轻量，适合中小型管理系统和演示型平台。",
            "数据库：SQLite，免部署、易迁移，适合课程设计、单机演示和小型数据平台原型。",
            "前端模板：Jinja2 + HTML + CSS + JavaScript，页面结构简单直观。",
            "图表展示：ECharts 负责 Web 端动态图表，Matplotlib 负责后端图片输出。",
            "预测模型：线性回归、复合增长率、近期趋势拟合和融合预测，适合小样本年度统计数据。",
            "解释分析：OpenAI 兼容大模型接口 + 本地规则兜底，兼顾智能解释和离线演示。",
            "数据导出：csv 模块生成 UTF-8 BOM CSV，保证中文在常见表格软件中正常打开。",
        ],
    )

    add_heading(doc, "八．异常处理与数据质量保障", 1)
    add_bullets(
        doc,
        [
            "表单输入使用 validators.py 校验年份、非负数、必填文本和人口数据字段完整性。",
            "采集模块支持 utf-8、gb18030 等编码回退，降低公开网页中文编码异常影响。",
            "人口数据新增和修改捕获 UNIQUE 冲突，提示同一地区同一年份数据已存在。",
            "扩展指标值录入前检查 indicator_code 是否存在，避免无效外键数据。",
            "预测模块要求至少 3 年历史数据，预测年数限制在 1-10 年，避免无依据外推。",
            "大模型接口调用失败时自动回退到本地规则解释，保证演示连续性。",
            "所有关键写操作写入 operation_logs，便于追踪问题和展示系统审计能力。",
        ],
    )

    add_heading(doc, "九．扩展预留设计", 1)
    add_bullets(
        doc,
        [
            "可在 collection_service 基础上扩展真实统计局接口定时采集和定时同步任务。",
            "可在 prediction_service 中增加 ARIMA、指数平滑、Prophet 等时间序列模型，并保留模型对比结果。",
            "可将 analysis_reports 扩展为 Word/PDF 报告导出，实现从分析到正式报告的完整输出。",
            "可增加用户管理、密码修改、角色权限字典和更细粒度菜单权限。",
            "可扩展地区层级到省、市、区县，并通过 parent_region 建立行政区划树。",
            "可将 SQLite 替换为 MySQL 或 PostgreSQL，以支持多人并发和更大规模数据。",
            "可将人口预警规则配置化，支持按指标阈值、地区类型和年份区间自定义预警策略。",
        ],
    )

    add_heading(doc, "十．总结", 1)
    add_paragraph(
        doc,
        "Population-Insight 的整体开发思路是以人口年度核心数据为中心，通过数据采集、数据管理、统计分析、可视化展示、机器学习预测、智能风险评估、报告保存和 CSV 导出构成完整业务闭环。项目在结构上采用配置层、数据层、服务层、Web 层、前端可视化层和输出文档层分层设计；在功能上覆盖课程设计常见的数据库建模、CRUD、查询筛选、统计排名、图表分析、预测算法、权限验证和成果材料整理要求，具有较好的演示完整性和后续扩展空间。"
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
