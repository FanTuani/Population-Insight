from __future__ import annotations

from datetime import datetime
from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from population_insight.db.connection import fetch_all, fetch_one


OUT_DIR = ROOT / "output" / "table_evidence"
DOC_PATH = ROOT / "output" / "table_crud_analysis.docx"
CONNECTION_SCREENSHOT = (
    Path.home()
    / "Pictures"
    / "Screenshots"
    / "屏幕截图 2026-06-07 185343.png"
)


def read_lines(path: str, start: int, end: int) -> str:
    lines = (ROOT / path).read_text(encoding="utf-8").splitlines()
    return "\n".join(
        f"{line_number + 1}: {lines[line_number]}"
        for line_number in range(start - 1, min(end, len(lines)))
    )


SNIPPETS = {
    "connection": read_lines("population_insight/db/connection.py", 75, 130),
    "users_read": read_lines("population_insight/services/auth_service.py", 1, 24),
    "users_page": read_lines("app.py", 665, 671),
    "seed_users": read_lines("population_insight/db/initializer.py", 406, 418),
    "population_insert": read_lines("population_insight/services/population_service.py", 70, 103),
    "population_query": read_lines("population_insight/services/population_service.py", 105, 142),
    "population_update_delete": read_lines("population_insight/services/population_service.py", 144, 181),
    "population_routes": read_lines("app.py", 225, 407),
    "log_service": read_lines("population_insight/services/log_service.py", 1, 28),
    "logs_route": read_lines("app.py", 652, 656),
    "regions_service": read_lines("population_insight/services/extension_service.py", 28, 53),
    "regions_route": read_lines("app.py", 684, 696),
    "sources_service": read_lines("population_insight/services/extension_service.py", 56, 92),
    "sources_route": read_lines("app.py", 699, 711),
    "indicators_service": read_lines("population_insight/services/extension_service.py", 95, 114),
    "indicators_route": read_lines("app.py", 714, 732),
    "indicator_values_service": read_lines("population_insight/services/extension_service.py", 117, 158),
    "indicator_values_route": read_lines("app.py", 735, 743),
    "reports_service": read_lines("population_insight/services/extension_service.py", 161, 177),
    "reports_route": read_lines("app.py", 746, 758),
    "national_seed": read_lines("population_insight/db/initializer.py", 512, 552),
    "national_service": read_lines("population_insight/services/national_series_service.py", 251, 358),
    "national_route": read_lines("app.py", 889, 896),
}


TABLES = [
    {
        "name": "users",
        "title": "用户表",
        "purpose": "保存系统登录账号、密码哈希、角色和创建时间，用于后台登录认证和权限判断。",
        "visual": "用户管理页面 /users 可查看用户列表；登录页会通过该表校验账号。",
        "crud": {
            "增": "已实现。系统初始化 seed_users() 写入默认 admin/viewer 账号；当前系统未开放前端新增用户表单。",
            "查": "已实现。authenticate_user() 查询登录用户，users_view() 在 /users 页面查询并展示用户列表。",
            "改": "未开放。当前系统没有修改用户角色或密码的页面/服务函数。",
            "删": "未开放。当前系统没有删除用户的页面/服务函数。",
        },
        "snippets": [("初始化新增默认用户", "seed_users"), ("登录与用户查询", "users_read"), ("用户页面路由", "users_page")],
    },
    {
        "name": "population_data",
        "title": "人口数据表",
        "purpose": "核心业务表，保存地区、年份、总人口、性别人口、出生率、死亡率、自然增长率、老龄化率、城镇化率、来源和质量标记。",
        "visual": "人口数据管理 /records 支持查询；/records/new 支持新增；编辑和删除入口在记录列表中。",
        "crud": {
            "增": "已实现。add_population_record() 校验表单后 INSERT INTO population_data。",
            "查": "已实现。query_population_records() 按地区、年份和指标区间动态查询。",
            "改": "已实现。update_population_record() 根据 id 更新记录。",
            "删": "已实现。delete_population_record() 根据 id 删除记录。",
        },
        "snippets": [("新增函数", "population_insert"), ("查询函数", "population_query"), ("修改与删除函数", "population_update_delete"), ("页面路由", "population_routes")],
    },
    {
        "name": "operation_logs",
        "title": "操作日志表",
        "purpose": "保存用户对数据的新增、修改、删除、扩展信息维护等操作记录，用于审计追踪。",
        "visual": "操作日志页面 /logs 可查看最近操作。",
        "crud": {
            "增": "已实现。log_operation() 在业务操作后自动 INSERT INTO operation_logs。",
            "查": "已实现。list_operation_logs() 查询日志，/logs 页面展示。",
            "改": "未开放。日志作为审计记录，不提供修改功能。",
            "删": "未开放。日志作为审计记录，不提供前端删除功能。",
        },
        "snippets": [("日志新增与查询函数", "log_service"), ("日志页面路由", "logs_route")],
    },
    {
        "name": "regions",
        "title": "地区档案表",
        "purpose": "保存地区名称、地区类型、行政区划代码、上级地区和备注，用于地区档案维护。",
        "visual": "地区档案页面 /regions 支持新增和查看。",
        "crud": {
            "增": "已实现。add_region() 写入 regions 表。",
            "查": "已实现。list_regions() 查询地区档案列表。",
            "改": "未开放。当前系统未提供地区档案编辑入口。",
            "删": "未开放。当前系统未提供地区档案删除入口。",
        },
        "snippets": [("地区新增与查询函数", "regions_service"), ("地区页面路由", "regions_route")],
    },
    {
        "name": "data_sources",
        "title": "数据来源表",
        "purpose": "保存数据来源名称、发布机构、来源链接、发布日期、可信等级和备注，用于数据质量说明。",
        "visual": "数据来源页面 /sources 支持新增和查看。",
        "crud": {
            "增": "已实现。add_data_source() 写入 data_sources 表。",
            "查": "已实现。list_data_sources() 查询数据来源列表。",
            "改": "未开放。当前系统未提供来源编辑入口。",
            "删": "未开放。当前系统未提供来源删除入口。",
        },
        "snippets": [("来源新增与查询函数", "sources_service"), ("来源页面路由", "sources_route")],
    },
    {
        "name": "population_indicators",
        "title": "扩展指标定义表",
        "purpose": "保存扩展指标编码、名称、单位和说明，为年度扩展指标值提供指标字典。",
        "visual": "扩展指标页面 /indicators 上半部分支持新增和查看指标定义。",
        "crud": {
            "增": "已实现。add_population_indicator() 写入 population_indicators 表。",
            "查": "已实现。list_population_indicators() 查询指标定义列表。",
            "改": "未开放。当前系统未提供指标定义编辑入口。",
            "删": "未开放。当前系统未提供指标定义删除入口。",
        },
        "snippets": [("指标定义新增与查询函数", "indicators_service"), ("指标页面路由", "indicators_route")],
    },
    {
        "name": "annual_indicator_values",
        "title": "年度扩展指标值表",
        "purpose": "保存某地区某年份某扩展指标的数值，和 population_indicators 通过 indicator_code 关联。",
        "visual": "扩展指标页面 /indicators 下半部分支持新增和查看年度指标值。",
        "crud": {
            "增": "已实现。add_annual_indicator_value() 写入 annual_indicator_values 表。",
            "查": "已实现。list_annual_indicator_values() 联表查询指标名称和单位。",
            "改": "未开放。当前系统未提供年度指标值编辑入口。",
            "删": "未开放。当前系统未提供年度指标值删除入口。",
        },
        "snippets": [("年度指标值新增与查询函数", "indicator_values_service"), ("年度指标值路由", "indicator_values_route")],
    },
    {
        "name": "analysis_reports",
        "title": "分析报告表",
        "purpose": "保存用户生成或填写的分析报告标题、筛选摘要、报告摘要和创建时间。",
        "visual": "分析报告页面 /reports 支持新增和查看；统计、预测等页面也可保存报告。",
        "crud": {
            "增": "已实现。add_analysis_report() 写入 analysis_reports 表。",
            "查": "已实现。list_analysis_reports() 查询报告列表。",
            "改": "未开放。当前系统未提供报告编辑入口。",
            "删": "未开放。当前系统未提供报告删除入口。",
        },
        "snippets": [("报告新增与查询函数", "reports_service"), ("报告页面路由", "reports_route")],
    },
    {
        "name": "national_population_series",
        "title": "全国长序列表",
        "purpose": "保存全国 1950-2025 年人口长序列，用于全国趋势图和长期预测。",
        "visual": "全国长序列页面 /national-series 展示摘要、趋势图和年度明细。",
        "crud": {
            "增": "已实现。seed_national_series() 初始化时插入全国长序列。",
            "查": "已实现。list_national_series()、get_national_trend_data() 查询并生成全国趋势图数据。",
            "改": "已实现于初始化同步。seed_national_series() 使用 ON DUPLICATE KEY UPDATE 更新已有年份数据；当前没有前端手动编辑入口。",
            "删": "未开放。全国长序列作为基础数据，不提供前端删除入口。",
        },
        "snippets": [("全国长序列初始化新增/更新", "national_seed"), ("全国长序列查询与趋势函数", "national_service"), ("全国趋势接口路由", "national_route")],
    },
]


def set_font(run, name: str = "宋体") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    set_font(run, "黑体")
    return paragraph


def para(doc: Document, text: str = ""):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def code_block(doc: Document, title: str, code: str) -> None:
    heading(doc, title, level=4)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(7.5)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    paragraph._p.get_or_add_pPr().append(shading)


def image_block(doc: Document, title: str, path: Path, max_width: float = 6.5) -> None:
    heading(doc, title, level=4)
    if not path.exists():
        para(doc, f"截图未找到：{path}")
        return
    with Image.open(path) as image:
        width = min(max_width, image.size[0] / 150)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption = doc.add_paragraph(f"截图文件：{path}")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        set_font(run)
        run.font.size = Pt(8.5)


def table_count(table: str) -> int:
    row = fetch_one(f"SELECT COUNT(*) AS count FROM {table}") or {"count": 0}
    return int(row["count"])


def describe_table(table: str) -> list[dict]:
    return fetch_all(f"DESCRIBE {table}")


def add_field_table(doc: Document, table: str) -> None:
    fields = describe_table(table)
    heading(doc, "字段结构", level=4)
    grid = doc.add_table(rows=1, cols=6)
    grid.style = "Table Grid"
    headers = ["字段", "类型", "可空", "键", "默认值", "额外信息"]
    for index, header in enumerate(headers):
        grid.rows[0].cells[index].text = header
    for field in fields:
        cells = grid.add_row().cells
        values = [
            field.get("Field", ""),
            field.get("Type", ""),
            field.get("Null", ""),
            field.get("Key", ""),
            field.get("Default", ""),
            field.get("Extra", ""),
        ]
        for index, value in enumerate(values):
            cells[index].text = "" if value is None else str(value)


def add_crud_table(doc: Document, table_info: dict) -> None:
    heading(doc, "CRUD 覆盖情况", level=4)
    grid = doc.add_table(rows=1, cols=2)
    grid.style = "Table Grid"
    grid.rows[0].cells[0].text = "操作"
    grid.rows[0].cells[1].text = "系统实现情况"
    for operation in ["增", "查", "改", "删"]:
        cells = grid.add_row().cells
        cells[0].text = operation
        cells[1].text = table_info["crud"][operation]


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MySQL 数据表增删改查设计与系统展示证明")
    run.bold = True
    run.font.size = Pt(18)
    set_font(run, "黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"项目：Population Insight    生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}")
    set_font(run)

    heading(doc, "一、MySQL 连接与数据库总览", level=2)
    para(doc, "本系统默认连接 MySQL 数据库 population_insight，使用 PyMySQL 完成连接、参数化查询、新增、修改和删除。")
    image_block(doc, "MySQL 连接成功截图", CONNECTION_SCREENSHOT)
    code_block(doc, "MySQL 连接与通用执行函数", SNIPPETS["connection"])

    heading(doc, "数据库表清单", level=3)
    overview = doc.add_table(rows=1, cols=5)
    overview.style = "Table Grid"
    for index, header in enumerate(["表名", "中文说明", "记录数", "可视页面", "CRUD 摘要"]):
        overview.rows[0].cells[index].text = header
    for table_info in TABLES:
        cells = overview.add_row().cells
        cells[0].text = table_info["name"]
        cells[1].text = table_info["title"]
        cells[2].text = str(table_count(table_info["name"]))
        cells[3].text = table_info["visual"]
        cells[4].text = "；".join(f"{key}:{value.split('。')[0]}" for key, value in table_info["crud"].items())

    heading(doc, "二、逐表增删改查分析", level=2)
    for index, table_info in enumerate(TABLES, start=1):
        name = table_info["name"]
        heading(doc, f"{index}. {name}（{table_info['title']}）", level=3)
        para(doc, f"表用途：{table_info['purpose']}")
        para(doc, f"系统展示：{table_info['visual']}")
        add_crud_table(doc, table_info)
        add_field_table(doc, name)

        page_shot = OUT_DIR / f"page_{name}.png"
        db_shot = OUT_DIR / f"db_{name}.png"
        image_block(doc, f"{name} 系统页面截图", page_shot)
        image_block(doc, f"{name} MySQL 表数据截图", db_shot)

        heading(doc, "对应代码片段", level=4)
        for title, snippet_key in table_info["snippets"]:
            code_block(doc, title, SNIPPETS[snippet_key])

        if index != len(TABLES):
            doc.add_page_break()

    heading(doc, "三、结论", level=2)
    para(doc, "系统已成功连接 MySQL，并围绕 9 张业务表完成数据初始化、查询展示、核心业务新增、部分表修改与删除。")
    para(doc, "其中 population_data 为核心业务表，系统提供完整新增、查询、修改、删除；其他配置类、日志类、报告类和长序列表根据业务需要提供新增/查询或初始化同步能力。未开放修改/删除入口的表，文档中已明确标注，避免把未实现功能误写为已实现。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Population Insight - MySQL 数据表 CRUD 分析"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_font(run)
        run.font.size = Pt(8.5)

    doc.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    build_doc()
