from __future__ import annotations

from datetime import datetime
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from population_insight.db.connection import fetch_one
OUT_DIR = ROOT / "output" / "mysql_evidence"
OUT_DIR.mkdir(parents=True, exist_ok=True)

CONNECTION_SCREENSHOT = (
    Path.home()
    / "Pictures"
    / "Screenshots"
    / "屏幕截图 2026-06-07 185343.png"
)
ADD_FORM_SCREENSHOT = OUT_DIR / "02_platform_add_form_filled.png"
AFTER_INSERT_SCREENSHOT = OUT_DIR / "03_platform_after_insert.png"
QUERY_RESULT_SCREENSHOT = OUT_DIR / "04_platform_query_result.png"
DOC_PATH = OUT_DIR / "mysql_basic_connection_insert_query_cn.docx"


def read_lines(path: str, start: int, end: int) -> str:
    lines = (ROOT / path).read_text(encoding="utf-8").splitlines()
    return "\n".join(
        f"{line_number + 1}: {lines[line_number]}"
        for line_number in range(start - 1, min(end, len(lines)))
    )


def set_font(run, font_name: str = "宋体") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    set_font(run, "黑体")
    return paragraph


def paragraph(doc: Document, text: str = ""):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run)
    return para


def code_block(doc: Document, title: str, code: str) -> None:
    heading(doc, title, level=3)
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    para._p.get_or_add_pPr().append(shading)


def image_block(doc: Document, title: str, path: Path, max_width: float = 6.3) -> None:
    heading(doc, title, level=3)
    if not path.exists():
        paragraph(doc, f"图片未找到：{path}")
        return

    with Image.open(path) as image:
        width = min(max_width, image.size[0] / 150)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Inches(width))

    caption = doc.add_paragraph(f"截图文件：{path}")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        set_font(run)
        run.font.size = Pt(9)


def latest_test_row() -> dict | None:
    return fetch_one(
        """
        SELECT id, region, year, total_population, male_population, female_population,
               birth_rate, death_rate, natural_growth_rate, data_quality
        FROM population_data
        WHERE data_quality = ?
        ORDER BY id DESC
        LIMIT 1
        """,
        ("test_insert_query",),
    )


def build_document() -> None:
    snippets = {
        "config": read_lines("population_insight/config.py", 10, 18),
        "connection": read_lines("population_insight/db/connection.py", 75, 130),
        "insert": read_lines("population_insight/services/population_service.py", 70, 103),
        "query": read_lines("population_insight/services/population_service.py", 105, 142),
        "routes": read_lines("app.py", 225, 366),
    }

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("3. MySQL 基础连接 + 增查功能证明材料")
    run.bold = True
    run.font.size = Pt(18)
    set_font(run, "黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"项目：Population Insight    生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}"
    )
    set_font(run)

    heading(doc, "一、评分点对应说明", level=2)
    for item in [
        "连接成功：JetBrains Data Sources and Drivers 测试 MySQL 连接，截图显示 Succeeded，DBMS 为 MySQL 8.0.31。",
        "新增功能：Web 平台“新增人口记录”页面提交测试记录，后端 add_population_record() 写入 MySQL。",
        "查询功能：Web 平台“人口数据管理”页面按地区和年份筛选，后端 query_population_records() 从 MySQL 查询并展示。",
        "数据库配置：项目默认数据库引擎已设为 MySQL，连接参数写入 population_insight/config.py。",
    ]:
        paragraph(doc, "• " + item)

    heading(doc, "二、MySQL 连接成功截图", level=2)
    image_block(doc, "图1：MySQL Data Source 测试连接成功", CONNECTION_SCREENSHOT)
    code_block(doc, "连接配置代码：population_insight/config.py", snippets["config"])
    code_block(doc, "连接与通用查询函数：population_insight/db/connection.py", snippets["connection"])

    heading(doc, "三、新增功能证明", level=2)
    paragraph(
        doc,
        "平台入口：登录系统后进入“人口数据管理 / 新增人口记录”。填写测试地区、年份、人口数据后提交，"
        "系统将数据写入 MySQL 的 population_data 表。",
    )
    image_block(doc, "图2：平台新增记录表单截图", ADD_FORM_SCREENSHOT)
    image_block(doc, "图3：提交新增后返回列表截图", AFTER_INSERT_SCREENSHOT)
    code_block(doc, "新增函数：population_insight/services/population_service.py", snippets["insert"])

    heading(doc, "四、查询功能证明", level=2)
    paragraph(
        doc,
        "平台入口：人口数据管理页面使用地区和年份筛选条件查询。截图中显示刚刚新增的测试记录，"
        "说明查询功能可从 MySQL 正常读取数据并展示到前端。",
    )
    image_block(doc, "图4：平台查询结果截图", QUERY_RESULT_SCREENSHOT)
    code_block(doc, "查询函数：population_insight/services/population_service.py", snippets["query"])
    code_block(doc, "Web 路由：app.py /records 与 /records/new", snippets["routes"])

    heading(doc, "五、MySQL 中测试记录核验", level=2)
    row = latest_test_row()
    if row:
        paragraph(doc, "通过项目 MySQL 连接层读取到的最新测试记录如下：")
        table = doc.add_table(rows=1, cols=len(row))
        table.style = "Table Grid"
        for index, key in enumerate(row.keys()):
            table.rows[0].cells[index].text = str(key)
        cells = table.add_row().cells
        for index, value in enumerate(row.values()):
            cells[index].text = str(value)
    else:
        paragraph(doc, "未读取到 data_quality=test_insert_query 的测试记录，请重新执行新增截图步骤。")

    heading(doc, "六、建议手工截图位置", level=2)
    for item in [
        "连接成功截图：JetBrains Data Sources and Drivers 窗口，保留左下角 Succeeded、DBMS、Ping 信息。",
        "新增功能截图：/records/new 页面，表单填好但提交前截图，能看到地区、年份、人口字段。",
        "查询展示截图：/records 页面按测试地区和年份筛选后的表格，能看到新增记录。",
        "代码截图：connection.py 中 get_connection/fetch_all/execute_write，population_service.py 中 add_population_record/query_population_records。",
    ]:
        paragraph(doc, "• " + item)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Population Insight - MySQL 基础连接 + 增查功能证明"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_font(run)
        run.font.size = Pt(9)

    doc.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    build_document()
